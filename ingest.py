import re
import uuid
import logging
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional
 
import pymupdf
import requests
from bs4 import BeautifulSoup
from docx import Document as DocxFile
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, VideoUnavailable
 
from readability import Document as ReadabilityDoc
READABILITY = True

@dataclass
class Document:
    text: str
    source: str
    source_type: str                      
    title: str = ""
    created_at: str = ""
    doc_id: str = ""
 
    def __post_init__(self):
        if not self.created_at:
            self.created_at = datetime.now(timezone.utc).isoformat()
        if not self.doc_id:
            self.doc_id = str(uuid.uuid5(uuid.NAMESPACE_URL, f"{self.source}{self.text[:200]}"))
 
    def preview(self, chars=300):
        return self.text[:chars].strip() + ("..." if len(self.text) > chars else "")
 
    def __repr__(self):
        return f"Document(id={self.doc_id!r}, type={self.source_type!r}, title={self.title!r}, words={len(self.text.split())})"
    
def _clean(text: str) -> str:
    lines = [l.strip() for l in text.splitlines() if len(l.strip()) > 1]
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()

def extract_pdf(path: str) -> Document:
    p = Path(path)
    doc = pymupdf.open(str(p))
    meta = doc.metadata or {}
    pages = []
 
    for i, page in enumerate(doc, 1):
        blocks = page.get_text("blocks", sort=True)
        text_chunks = [b[4].strip() for b in blocks if b[6] == 0 and b[4].strip()]
        if text_chunks:
            pages.append(f"[Page {i}]\n" + "\n".join(text_chunks))
 
    doc.close()
    return Document(
        text=_clean("\n\n".join(pages)),
        source=str(p.resolve()),
        source_type="pdf",
        title=meta.get("title", p.stem).strip() or p.stem,
    )
 
def extract_docx(path: str) -> Document:
    p = Path(path)
    doc = DocxFile(str(p))
 
    heading_map = {"Heading 1": "# ", "Heading 2": "## ", "Heading 3": "### ", "Title": "# "}
    lines = []
 
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(heading_map.get(para.style.name, "") + t)
 
    for table in doc.tables:
        lines.append("[Table]")
        for row in table.rows:
            lines.append(" | ".join(c.text.strip() for c in row.cells))
 
    props = doc.core_properties
    return Document(
        text=_clean("\n".join(lines)), source=str(p.resolve()), source_type="docx",
        title=props.title or p.stem, author=props.author or "",
    )

def extract_youtube(url: str) -> Document:
    video_id = url.split("v=")[-1].split("&")[0]
    ytt_api = YouTubeTranscriptApi()

    try:
        transcript = ytt_api.fetch(video_id)
    except TranscriptsDisabled:
        raise ValueError(f"Transcripts disabled for video: {video_id}")
    except VideoUnavailable:
        raise ValueError(f"Video unavailable: {video_id}")
 
    text = " ".join([entry.text for entry in transcript])
    title = video_id

    try:
        resp = requests.get(f"https://www.youtube.com/watch?v={video_id}", timeout=10)
        m = re.search(r"<title>(.+?) - YouTube</title>", resp.text)
        if m:
            title = m.group(1).strip()
    except Exception:
        pass
 
    return Document(
        text=_clean(text),
        source=url,
        source_type="youtube",
        title=title,
    )

def extract_url(url: str) -> Document:
    resp = requests.get(url, timeout=15, headers={
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/146.0.0.0 Safari/537.36"
        })
    resp.raise_for_status()
 
    doc = ReadabilityDoc(resp.text)
    title = doc.title()
    soup = BeautifulSoup(doc.summary(), "html.parser")
    text = _clean(soup.get_text(separator="\n"))
 
    return Document(
        text=text,
        source=url,
        source_type="url",
        title=title,
    )

def extract_note(text: str, title: str = "") -> Document:
    return Document(
        text=_clean(text),
        source="note",
        source_type="note",
        title=title or text[:50].strip(),
    )